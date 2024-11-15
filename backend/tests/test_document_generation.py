from docx import Document as DocxDocument
from pipeline.pipeline import Document
import os

def test_document_generation():
    # Create a mock processed document with sample sections and tables
    mock_document = Document(
        sections={
            "Water": "This is a sample water section discussing floods and ports.",
            "Fire": "This is a sample fire section about wildfires and stations.",
            "Administrative": "This discusses administrative matters and employees.",
            "Other": "Miscellaneous information goes here."
        },
        tables=[
            {
                'headers': ['Name', 'Position', 'Department'],
                'rows': [
                    {'Name': 'John Doe', 'Position': 'Manager', 'Department': 'Sales'},
                    {'Name': 'Jane Smith', 'Position': 'Engineer', 'Department': 'IT'},
                ]
            },
            {
                'headers': ['Item', 'Quantity', 'Cost'],
                'rows': [
                    {'Item': 'Laptop', 'Quantity': '5', 'Cost': '$5000'},
                    {'Item': 'Printer', 'Quantity': '2', 'Cost': '$1000'},
                ]
            }
        ]
    )

    # Create Word document
    doc = DocxDocument()
    doc.add_heading('Document Summary', 0)
    doc.add_paragraph('Summary Type: Test Summary')
    
    # Add document information
    doc.add_heading('Document: test_document.pdf', level=1)
    doc.add_paragraph('Type: Test Document')
    
    # Add each section to the Word document
    for section_name, content in mock_document.sections.items():
        doc.add_heading(section_name, level=2)
        doc.add_paragraph(content)
    
    # Add tables
    doc.add_heading('Extracted Tables', level=1)
    for i, table_data in enumerate(mock_document.tables):
        doc.add_heading(f'Table {i + 1}', level=2)
        
        # Create table in Word document
        table = doc.add_table(rows=1, cols=len(table_data['headers']))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(table_data['headers']):
            header_cells[i].text = header
        
        # Add data rows
        for row_data in table_data['rows']:
            row_cells = table.add_row().cells
            for i, header in enumerate(table_data['headers']):
                row_cells[i].text = row_data[header]
    
    # Save the generated summary
    os.makedirs('test_output', exist_ok=True)
    output_path = 'test_output/test_summary.docx'
    doc.save(output_path)
    print(f"Test document generated at: {output_path}")

if __name__ == "__main__":
    test_document_generation()
