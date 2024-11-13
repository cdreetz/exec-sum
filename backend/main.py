import os
import asyncio
from typing import List
from docx import Document as DocxDocument
from config import settings
from openai import AzureOpenAI
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from pipeline.template import EXAMPLE_TEMPLATES
from pipeline.pipeline import DocumentProcessor, Document
from pipeline.document import Document as ProcessedDocument

app = FastAPI()

origins = ["http://localhost:3000"]  # Adjust this to match your frontend URL

# Add CORS middleware and static file mounting with error handling
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.CORS_ORIGINS,
    allow_credentials=settings.CORS_ALLOW_CREDENTIALS,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Check if directories exist before mounting
if os.path.exists(settings.STATIC_DIR):
    app.mount("/static", StaticFiles(directory=settings.STATIC_DIR), name="static")
if os.path.exists(settings.ASSETS_DIR):
    app.mount("/assets", StaticFiles(directory=settings.ASSETS_DIR), name="assets")

# Initialize Azure clients
doc_client = DocumentIntelligenceClient(
    endpoint=settings.AZURE_ENDPOINT,
    credential=AzureKeyCredential(settings.AZURE_API_KEY)
)

openai_client = AzureOpenAI(
    api_version="2024-08-01-preview",
    azure_endpoint=settings.OPENAI_ENDPOINT,
    api_key=settings.AZURE_OPENAI_API_KEY
)

@app.get("/{full_path:path}")
async def serve_react_app(full_path: str):
    file_path = os.path.join(settings.BUILD_DIR, full_path)
    if os.path.isfile(file_path):
        return FileResponse(file_path)

    return FileResponse(os.path.join(settings.BUILD_DIR, "index.html"))

@app.post("/example_generate_summary")
async def example_generate_summary(
    file: UploadFile = File(...),
    type: str = Form(...),
    summary_type: str = Form(...)):
    
    # Create a temporary directory for uploaded files
    temp_dir = "temp_uploads"
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        # Save uploaded file
        file_location = os.path.join(temp_dir, file.filename)
        with open(file_location, "wb+") as file_object:
            content = await file.read()
            file_object.write(content)

        # Simulate processing time
        await asyncio.sleep(2)

        # Create a Word document with summary
        doc = Document()
        doc.add_heading(f'Document Summary', 0)
        doc.add_paragraph(f'Summary Type: {summary_type}')
        
        # Add document information
        doc.add_heading(f'Document: {file.filename}', level=1)
        doc.add_paragraph(f'Type: {type}')
        doc.add_paragraph('This is a sample summary for the document.')
            
        # Save the generated summary with .docx extension
        output_path = os.path.join(temp_dir, "generated_summary.docx")
        doc.save(output_path)
        
        return FileResponse(
            path=output_path,
            filename="generated_summary.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        # Cleanup temporary files
        if os.path.exists(file_location):
            os.remove(file_location)

@app.post("/generate_summary")
async def generate_summary(
    file: UploadFile = File(...),
    type: str = Form(...),
    summary_type: str = Form(...)):
    
    temp_dir = "temp_uploads"
    os.makedirs(temp_dir, exist_ok=True)
    file_location = None
    
    try:
        # Save uploaded file
        file_location = os.path.join(temp_dir, file.filename)
        with open(file_location, "wb+") as file_object:
            content = await file.read()
            file_object.write(content)

        # Initialize pipeline with example template based on summary type
        pipeline = DocumentProcessor(doc_client, openai_client)
        example_document = EXAMPLE_TEMPLATES.get(
            summary_type.lower(), 
            EXAMPLE_TEMPLATES["brief"]
        )

        # Process the document and create a new Document instance
        processed_content = pipeline.process_document(
            file_location, 
            example_document
        )
        
        # Create Word document with processed content
        doc = DocxDocument()
        doc.add_heading(f'Document Summary', 0)
        doc.add_paragraph(f'Summary Type: {summary_type}')
        
        # Add document information
        doc.add_heading(f'Document: {file.filename}', level=1)
        doc.add_paragraph(f'Type: {type}')
        
        # Add each section to the Word document
        for section_name, content in processed_content.sections.items():
            doc.add_heading(section_name, level=2)
            doc.add_paragraph(content)
            
        # Save the generated summary
        output_path = os.path.join(temp_dir, "generated_summary.docx")
        doc.save(output_path)
        
        return FileResponse(
            path=output_path,
            filename="generated_summary.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        # Cleanup temporary files
        if file_location and os.path.exists(file_location):
            try:
                os.remove(file_location)
            except Exception as e:
                print(f"Error cleaning up file: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
